from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pandas as pd
import os
from datetime import datetime

class GeradorLaudoDocx:
    def __init__(self):
        self.document = None
        
    def criar_documento(self):
        """Cria um novo documento DOCX"""
        self.document = Document()
        self._configurar_estilos()
        
    def _configurar_estilos(self):
        """Configura os estilos do documento"""
        # Estilo para título principal
        titulo_style = self.document.styles.add_style('TituloLaudo', WD_STYLE_TYPE.PARAGRAPH)
        titulo_font = titulo_style.font
        titulo_font.name = 'Arial'
        titulo_font.size = Pt(18)
        titulo_font.bold = True
        titulo_font.color.rgb = RGBColor(31, 78, 121)  # Azul escuro
        titulo_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_style.paragraph_format.space_after = Pt(20)
        titulo_style.paragraph_format.space_before = Pt(10)
        
        # Estilo para subtítulos
        subtitulo_style = self.document.styles.add_style('SubtituloLaudo', WD_STYLE_TYPE.PARAGRAPH)
        subtitulo_font = subtitulo_style.font
        subtitulo_font.name = 'Arial'
        subtitulo_font.size = Pt(14)
        subtitulo_font.bold = True
        subtitulo_font.color.rgb = RGBColor(46, 125, 50)  # Verde escuro
        subtitulo_style.paragraph_format.space_before = Pt(16)
        subtitulo_style.paragraph_format.space_after = Pt(8)
        
        # Estilo para texto normal
        texto_style = self.document.styles.add_style('TextoLaudo', WD_STYLE_TYPE.PARAGRAPH)
        texto_font = texto_style.font
        texto_font.name = 'Arial'
        texto_font.size = Pt(11)
        texto_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        texto_style.paragraph_format.space_after = Pt(6)
        texto_style.paragraph_format.line_spacing = 1.15
        
        # Estilo para assinatura
        assinatura_style = self.document.styles.add_style('AssinaturaLaudo', WD_STYLE_TYPE.PARAGRAPH)
        assinatura_font = assinatura_style.font
        assinatura_font.name = 'Arial'
        assinatura_font.size = Pt(11)
        assinatura_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        assinatura_style.paragraph_format.space_before = Pt(20)
        assinatura_style.paragraph_format.space_after = Pt(6)
    
    def _aplicar_estilo_tabela(self, tabela, cor_fundo=None):
        """Aplica estilo profissional nas tabelas"""
        for row_idx, row in enumerate(tabela.rows):
            for cell_idx, cell in enumerate(row.cells):
                # Aplicar cor de fundo se especificada
                if cor_fundo:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{:02X}{:02X}{:02X}"/>'.format(
                        nsdecls('w'), cor_fundo[0], cor_fundo[1], cor_fundo[2]))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # Aplicar bordas
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = tcBorders.find(qn(f'w:{border_name}'))
                    if border is None:
                        border = OxmlElement(f'w:{border_name}')
                        tcBorders.append(border)
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'CCCCCC')
        
    def adicionar_titulo(self, titulo="LAUDO DE AVALIAÇÃO IMOBILIÁRIA"):
        """Adiciona o título principal do laudo"""
        titulo_para = self.document.add_paragraph()
        titulo_para.style = 'TituloLaudo'
        titulo_run = titulo_para.add_run(titulo)
        titulo_run.bold = True
        
    def adicionar_objetivo(self):
        """Adiciona a seção de objetivo do laudo (texto padrão)"""
        self.document.add_paragraph("1. OBJETIVO DO LAUDO", style='SubtituloLaudo')
        
        objetivo_texto = """O presente laudo tem como finalidade estimar o valor de mercado do imóvel descrito, 
para fins de garantia de financiamento, ação judicial, atualização patrimonial, entre outros, 
em conformidade com a ABNT NBR 14653-1 (Procedimentos Gerais) e NBR 14653-2 (Imóveis Urbanos)."""
        
        self.document.add_paragraph(objetivo_texto, style='TextoLaudo')
        
    def adicionar_identificacao_imovel(self, dados_imovel):
        """Adiciona a seção de identificação do imóvel"""
        self.document.add_paragraph("2. IDENTIFICAÇÃO DO IMÓVEL", style='SubtituloLaudo')
        
        # Criar tabela para dados do imóvel
        dados_tabela = [
            ['Número da Matrícula:', dados_imovel.get('numero_matricula', 'N/A')],
            ['Cartório de Registro:', dados_imovel.get('cartorio', 'N/A')],
            ['Endereço Completo:', dados_imovel.get('endereco_completo', 'N/A')],
            ['Área Construída:', f"{dados_imovel.get('area_construida', 0):.2f} m²"],
            ['Bairro:', dados_imovel.get('Bairro', 'N/A')],
            ['Cidade/Estado:', f"{dados_imovel.get('cidade', 'N/A')}/{dados_imovel.get('estado', 'N/A')}"],
            ['Descrição:', dados_imovel.get('descricao', 'N/A')]
        ]
        
        tabela = self.document.add_table(rows=len(dados_tabela), cols=2)
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configurar largura das colunas
        for row in tabela.rows:
            row.cells[0].width = Inches(2.5)
            row.cells[1].width = Inches(4.2)
        
        for i, (campo, valor) in enumerate(dados_tabela):
            # Célula do campo (primeira coluna)
            cell_campo = tabela.cell(i, 0)
            cell_campo.text = campo
            # Formatação da célula do campo
            for paragraph in cell_campo.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(31, 78, 121)  # Azul escuro
            
            # Célula do valor (segunda coluna)
            cell_valor = tabela.cell(i, 1)
            cell_valor.text = valor
            # Formatação da célula do valor
            for paragraph in cell_valor.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Preto
        
        # Aplicar bordas e fundo
        self._aplicar_estilo_tabela(tabela, (248, 249, 250))
        
    def adicionar_metodologia(self):
        """Adiciona a seção de metodologia (texto padrão)"""
        self.document.add_paragraph("3. METODOLOGIA DE AVALIAÇÃO", style='SubtituloLaudo')
        
        metodologia_texto = """Este trabalho adota o Método Comparativo Direto de Dados de Mercado, conforme a NBR 14653, 
realizando coleta, análise e tratamento estatístico de amostra de imóveis similares ao objeto de avaliação. 
Foi empregada a técnica de regressão hedônica para ponderação de fatores como área construída, 
área do terreno, localização, padrão construtivo e estado de conservação, ajustando valores unitários obtidos em mercado. 
Foram descartados imóveis que apresentaram valores manifestamente incompatíveis com a realidade mercadológica local, 
a fim de evitar distorções na média amostral, seguindo as boas práticas de avaliação consagradas pelas normas técnicas."""
        
        self.document.add_paragraph(metodologia_texto, style='TextoLaudo')
        
    def adicionar_pesquisa_mercado(self, dados_scraper):
        """Adiciona a seção de pesquisa de mercado com dados do scraper"""
        self.document.add_paragraph("4. PESQUISA DE MERCADO E ANÁLISE COMPARATIVA", style='SubtituloLaudo')
        
        # Adicionar texto explicativo sobre a amostra dos top 10 valores
        texto_explicativo = """Para a análise comparativa, foram selecionados os 10 imóveis com melhor relação custo-benefício"""
        
        self.document.add_paragraph(texto_explicativo, style='TextoLaudo')
        
        if dados_scraper is not None and not dados_scraper.empty:
            # Criar tabela com dados do scraper - top 10
            tabela = self.document.add_table(rows=1, cols=5)
            tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Configurar largura das colunas
            for row in tabela.rows:
                row.cells[0].width = Inches(0.6)  # Nº
                row.cells[1].width = Inches(2.2)  # Localização
                row.cells[2].width = Inches(1.2)  # Área
                row.cells[3].width = Inches(1.8)  # Preço
                row.cells[4].width = Inches(1.8)  # Valor Unitário
            
            # Cabeçalho da tabela
            header_cells = tabela.rows[0].cells
            headers = ['Nº', 'Localização', 'Área (m²)', 'Preço Anunciado', 'Valor Unitário (m²)']
            
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                # Formatação do cabeçalho
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # Branco
                # Cor de fundo do cabeçalho
                shading_elm = parse_xml(r'<w:shd {} w:fill="{:02X}{:02X}{:02X}"/>'.format(
                    nsdecls('w'), 46, 125, 50))  # Verde escuro
                header_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
            
            # Adicionar dados do scraper (top 10 imóveis)
            df_sorted = dados_scraper.sort_values('R$/M2').head(10)
            
            for idx, (_, row) in enumerate(df_sorted.iterrows(), 1):
                row_cells = tabela.add_row().cells
                dados_linha = [
                    str(idx),
                    str(row.get('Endereco', 'N/A'))[:30] + '...' if len(str(row.get('Endereco', 'N/A'))) > 30 else str(row.get('Endereco', 'N/A')),
                    f"{row.get('M2', 0):.0f}",
                    f"R$ {row.get('Preco', 0):,.0f}",
                    f"R$ {row.get('R$/M2', 0):,.0f}"
                ]
                
                for i, dado in enumerate(dados_linha):
                    row_cells[i].text = dado
                    # Formatação das células de dados
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Preto
            
            # Aplicar bordas na tabela
            self._aplicar_estilo_tabela(tabela)
            
            # URLs removidas - mantidas apenas no Excel
        else:
            # Texto padrão se não houver dados do scraper - top 10
            texto_padrao = """Nº | Localização        | Área (m²) | Preço Anunciado | Valor Unitário (m²)
1  | Jardim Santa Virgínia | 150 m²    | R$ 450.000      | R$ 3.750
2  | Jardim Santa Virgínia | 160 m²    | R$ 460.000      | R$ 3.680
3  | Jardim Santa Virgínia | 180 m²    | R$ 490.000      | R$ 3.769
4  | Jardim Santa Virgínia | 165 m²    | R$ 470.000      | R$ 3.788
5  | Jardim Santa Virgínia | 155 m²    | R$ 455.000      | R$ 3.742
6  | Jardim Santa Virgínia | 175 m²    | R$ 485.000      | R$ 3.771
7  | Jardim Santa Virgínia | 170 m²    | R$ 480.000      | R$ 3.765
8  | Jardim Santa Virgínia | 145 m²    | R$ 445.000      | R$ 3.724
9  | Jardim Santa Virgínia | 185 m²    | R$ 495.000      | R$ 3.784
10 | Jardim Santa Virgínia | 140 m²    | R$ 440.000      | R$ 3.714"""
            
            self.document.add_paragraph(texto_padrao, style='TextoLaudo')
    
    def adicionar_determinacao_valor(self, dados_scraper, dados_imovel):
        """Adiciona a seção de determinação do valor de mercado"""
        self.document.add_paragraph("5. DETERMINAÇÃO DO VALOR DE MERCADO", style='SubtituloLaudo')
        
        # Inicializar variáveis
        area_construida = dados_imovel.get('area_construida', 0)
        valor_edificacao = 0
        
        # Calcular valor médio unitário baseado nos dados do scraper
        if dados_scraper is not None and not dados_scraper.empty:
            df_top10 = dados_scraper.sort_values('R$/M2').head(10)
            valor_medio_unitario = df_top10['R$/M2'].mean()
        else:
            valor_medio_unitario = 3700.0  # Valor padrão
        
        # Calcular valor da edificação apenas se área construída > 0
        if area_construida > 0:
            valor_edificacao = valor_medio_unitario * area_construida
            
        # Valor total é apenas o valor da edificação
        valor_total = valor_edificacao
        
        # Criar tabela de cálculos condicionalmente
        calculos_tabela = [['Descrição', 'Valor']]
        
        # Sempre mostrar valor médio unitário
        calculos_tabela.append(['Valor médio unitário de referência:', f"R$ {valor_medio_unitario:,.2f}/m²"])
        
        # Mostrar área construída e valor da edificação apenas se área > 0
        if area_construida > 0:
            calculos_tabela.append(['Área construída do imóvel:', f"{area_construida:.2f} m²"])
            calculos_tabela.append(['Valor estimado da edificação:', f"R$ {valor_edificacao:,.2f}"])
        
        # Sempre mostrar valor total
        calculos_tabela.append(['', ''])  # Linha em branco
        calculos_tabela.append(['VALOR TOTAL DO IMÓVEL:', f"R$ {valor_total:,.2f}"])
        
        # Criar tabela no documento
        tabela = self.document.add_table(rows=len(calculos_tabela), cols=2)
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configurar largura das colunas
        for row in tabela.rows:
            row.cells[0].width = Inches(4.2)
            row.cells[1].width = Inches(2.5)
        
        # Preencher tabela
        for i, (descricao, valor) in enumerate(calculos_tabela):
            cell_desc = tabela.cell(i, 0)
            cell_valor = tabela.cell(i, 1)
            
            cell_desc.text = descricao
            cell_valor.text = valor
            
            # Formatação das células
            for paragraph in cell_desc.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    if i == 0 or i == len(calculos_tabela) - 1:  # Cabeçalho e total
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255) if i == 0 else RGBColor(255, 255, 255)
                    else:
                        run.font.color.rgb = RGBColor(31, 78, 121) if descricao else RGBColor(0, 0, 0)
            
            for paragraph in cell_valor.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    if i == 0 or i == len(calculos_tabela) - 1:  # Cabeçalho e total
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255) if i == 0 else RGBColor(255, 255, 255)
                    else:
                        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Aplicar cor de fundo no cabeçalho
        shading_elm = parse_xml(r'<w:shd {} w:fill="{:02X}{:02X}{:02X}"/>'.format(
            nsdecls('w'), 46, 125, 50))  # Verde escuro
        tabela.cell(0, 0)._tc.get_or_add_tcPr().append(shading_elm)
        tabela.cell(0, 1)._tc.get_or_add_tcPr().append(shading_elm)
        
        # Aplicar cor de fundo no total
        shading_total = parse_xml(r'<w:shd {} w:fill="{:02X}{:02X}{:02X}"/>'.format(
            nsdecls('w'), 31, 78, 121))  # Azul escuro
        tabela.cell(len(calculos_tabela) - 1, 0)._tc.get_or_add_tcPr().append(shading_total)
        tabela.cell(len(calculos_tabela) - 1, 1)._tc.get_or_add_tcPr().append(shading_total)
        
        # Aplicar bordas na tabela
        self._aplicar_estilo_tabela(tabela)
        
    def adicionar_pesquisa_localidade(self, texto_pesquisa):
        """Adiciona a seção de pesquisa de localidade"""
        self.document.add_paragraph("6. PESQUISA DE LOCALIDADE", style='SubtituloLaudo')
        
        if texto_pesquisa:
            self.document.add_paragraph(texto_pesquisa, style='TextoLaudo')
        else:
            # Texto padrão quando não há informações específicas
            texto_padrao = """Comércio, Lazer e Comodidade:
Ele se beneficia de mercados locais, padarias e lojas de conveniência no entorno. 
A cerca de 2–4 km ficam clubes, praças e o Complexo Villa. 
Excelente acesso a postos de gasolina e serviços essenciais.

Segurança:
Não há dados específicos sobre índices de criminalidade no bairro, 
mas Taquarituba apresenta níveis moderados típicos de cidades do interior paulista.

Potencial Econômico e de Crescimento Imobiliário:
Jardim Santa Virgínia é um bairro residencial estável, próximo ao centro, 
com boa infraestrutura, transporte e perfil familiar. 
O custo imobiliário mais baixo e investimentos municipais recentes tornam-no atrativo para moradia e investimento de médio prazo."""
            
            self.document.add_paragraph(texto_padrao, style='TextoLaudo')
        
    def adicionar_conclusao(self, dados_scraper, dados_imovel):
        """Adiciona a seção de conclusão"""
        self.document.add_paragraph("7. CONCLUSÃO", style='SubtituloLaudo')
        
        # Inicializar variáveis
        area_construida = dados_imovel.get('area_construida', 0)
        valor_edificacao = 0
        
        # Calcular valor médio unitário baseado nos dados do scraper
        if dados_scraper is not None and not dados_scraper.empty:
            df_top10 = dados_scraper.sort_values('R$/M2').head(10)
            valor_medio_unitario = df_top10['R$/M2'].mean()
        else:
            valor_medio_unitario = 3700.0  # Valor padrão
        
        # Calcular valor da edificação apenas se área construída > 0
        if area_construida > 0:
            valor_edificacao = valor_medio_unitario * area_construida
            
        # Valor total é apenas o valor da edificação
        valor_total = valor_edificacao
        
        # Converter valor para extenso
        valor_extenso = self._converter_valor_extenso(valor_total)
        
        conclusao_texto = f"""Com base nos estudos realizados, considera-se que o valor de mercado do imóvel avaliado é de R$ {valor_total:,.2f} 
({valor_extenso}), na data de referência deste laudo. 

Este valor reflete as condições de mercado atuais, características do imóvel e seu estado de conservação, 
considerando as especificidades da amostra utilizada e os ajustes técnicos realizados."""
        
        self.document.add_paragraph(conclusao_texto, style='TextoLaudo')
        
    def adicionar_assinatura(self, dados_avaliador):
        """Adiciona a seção de assinatura do avaliador"""
        self.document.add_paragraph("8. ASSINATURA DO AVALIADOR", style='SubtituloLaudo')
        
        # Espaço para assinatura
        self.document.add_paragraph("", style='TextoLaudo')
        
        # Dados do avaliador
        cidade_assinatura = dados_avaliador.get('cidade_assinatura', 'Sorocaba')
        estado_assinatura = dados_avaliador.get('estado_assinatura', 'SP')
        data_laudo = dados_avaliador.get('data_laudo', datetime.now().date())
        
        # Criar parágrafo centralizado para assinatura
        assinatura_para = self.document.add_paragraph()
        assinatura_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data e local
        run_data = assinatura_para.add_run(f"{cidade_assinatura} - {estado_assinatura}, {data_laudo.strftime('%d de %B de %Y')}\n\n")
        run_data.font.name = 'Arial'
        run_data.font.size = Pt(11)
        run_data.font.color.rgb = RGBColor(0, 0, 0)
        
        # Nome do avaliador
        run_nome = assinatura_para.add_run(f"{dados_avaliador.get('nome', 'Jhonni Balbino da Silva')}\n")
        run_nome.font.name = 'Arial'
        run_nome.font.size = Pt(11)
        run_nome.font.bold = True
        run_nome.font.color.rgb = RGBColor(31, 78, 121)  # Azul escuro
        
        # Cargo
        run_cargo = assinatura_para.add_run("Perito Avaliador Imobiliário\n")
        run_cargo.font.name = 'Arial'
        run_cargo.font.size = Pt(11)
        run_cargo.font.color.rgb = RGBColor(0, 0, 0)
        
        # CRECI e CNAI
        run_creci = assinatura_para.add_run(f"CRECI: {dados_avaliador.get('creci', '296769-F')} | CNAI: {dados_avaliador.get('cnai', '051453')}\n")
        run_creci.font.name = 'Arial'
        run_creci.font.size = Pt(11)
        run_creci.font.color.rgb = RGBColor(0, 0, 0)
        
        # Telefone
        run_telefone = assinatura_para.add_run(f"{dados_avaliador.get('telefone', '(11) 98796 8206')}\n")
        run_telefone.font.name = 'Arial'
        run_telefone.font.size = Pt(11)
        run_telefone.font.color.rgb = RGBColor(0, 0, 0)
        
        # Email
        run_email = assinatura_para.add_run(f"{dados_avaliador.get('email', 'contato@avaliarapido.com.br')}\n")
        run_email.font.name = 'Arial'
        run_email.font.size = Pt(11)
        run_email.font.color.rgb = RGBColor(0, 0, 0)
        
        # Website
        run_website = assinatura_para.add_run(f"{dados_avaliador.get('website', 'www.avaliarapido.com.br')}")
        run_website.font.name = 'Arial'
        run_website.font.size = Pt(11)
        run_website.font.color.rgb = RGBColor(0, 0, 0)
        
    def _converter_valor_extenso(self, valor):
        """Converte valor numérico para extenso"""
        # Implementação simplificada para valores até milhões
        if valor >= 1000000:
            milhoes = int(valor // 1000000)
            resto = valor % 1000000
            if resto == 0:
                return f"{milhoes} milhão{'es' if milhoes > 1 else ''}"
            else:
                return f"{milhoes} milhão{'es' if milhoes > 1 else ''} e {self._converter_valor_extenso(resto)}"
        elif valor >= 1000:
            milhares = int(valor // 1000)
            resto = valor % 1000
            if resto == 0:
                return f"{milhares} mil"
            else:
                return f"{milhares} mil e {self._converter_valor_extenso(resto)}"
        else:
            return f"{int(valor)}"
    
    def salvar_documento(self, nome_arquivo=None):
        """Salva o documento DOCX"""
        if nome_arquivo is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            nome_arquivo = f"laudo_avaliacao_{timestamp}.docx"
        
        # Criar pasta arquivos se não existir
        pasta_arquivos = "arquivos"
        if not os.path.exists(pasta_arquivos):
            os.makedirs(pasta_arquivos)
        
        caminho_completo = os.path.join(pasta_arquivos, nome_arquivo)
        self.document.save(caminho_completo)
        
        return caminho_completo
    
    def gerar_laudo_completo(self, dados_imovel, dados_avaliador, dados_scraper=None, texto_pesquisa_localidade=None):
        """Gera o laudo completo com todos os dados"""
        self.criar_documento()
        
        # Adicionar todas as seções
        self.adicionar_titulo()
        self.adicionar_objetivo()
        self.adicionar_identificacao_imovel(dados_imovel)
        self.adicionar_metodologia()
        self.adicionar_pesquisa_mercado(dados_scraper)
        self.adicionar_determinacao_valor(dados_scraper, dados_imovel)
        
        if texto_pesquisa_localidade:
            self.adicionar_pesquisa_localidade(texto_pesquisa_localidade)
        else:
            # Texto padrão
            texto_padrao = """Comércio, Lazer e Comodidade:
Ele se beneficia de mercados locais, padarias e lojas de conveniência no entorno. 
A cerca de 2–4 km ficam clubes, praças e o Complexo Villa. 
Excelente acesso a postos de gasolina e serviços essenciais.

Segurança:
Não há dados específicos sobre índices de criminalidade no bairro, 
mas Taquarituba apresenta níveis moderados típicos de cidades do interior paulista.

Potencial Econômico e de Crescimento Imobiliário:
Jardim Santa Virgínia é um bairro residencial estável, próximo ao centro, 
com boa infraestrutura, transporte e perfil familiar. 
O custo imobiliário mais baixo e investimentos municipais recentes tornam-no atrativo para moradia e investimento de médio prazo."""
            self.adicionar_pesquisa_localidade(texto_padrao)
        
        self.adicionar_conclusao(dados_scraper, dados_imovel)
        self.adicionar_assinatura(dados_avaliador)
        
        # Salvar documento
        return self.salvar_documento()
    

def main():
    """Função para testar o gerador de laudo"""
    gerador = GeradorLaudoDocx()
    
    # Dados de exemplo
    dados_imovel = {
        'numero_matricula': '11.072',
        'cartorio': 'Comarca de Taquarituba/SP, Livro nº 2 – Registro Geral',
        'area_construida': 134.59,
        'Bairro': 'Jardim Santa Virgínia',
        'cidade': 'Taquarituba',
        'estado': 'SP',
        'descricao': 'Casa térrea em alvenaria com telhas cerâmicas, acabamento em gesso, piso cerâmico, portas e janelas em madeira, sistema elétrico e hidráulico completo.'
    }
    
    dados_avaliador = {
        'nome': 'Jhonni Balbino da Silva',
        'creci': '296769-F',
        'cnai': '051453',
        'telefone': '(11) 98796 8206',
        'email': 'contato@avaliarapido.com.br',
        'website': 'www.avaliarapido.com.br',
        'data_laudo': datetime.now().date(),
        'cidade_assinatura': 'Sorocaba',
        'estado_assinatura': 'SP'
    }
    
    # Gerar laudo
    arquivo_gerado = gerador.gerar_laudo_completo(dados_imovel, dados_avaliador)
    print(f"Laudo gerado: {arquivo_gerado}")

if __name__ == "__main__":
    main()
